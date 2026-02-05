# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: MIT-0

"""
Document converter for handling various document formats.

This module provides functionality to convert different document formats
(Plain Text, CSV, Excel, Word) into page images and text outputs
consistent with PDF processing.
"""

import io
import logging
import os
import tempfile
from typing import List, Tuple

from PIL import Image, ImageDraw, ImageFont

logger = logging.getLogger(__name__)


class DocumentConverter:
    """Converter for various document formats to images and text."""

    def __init__(self, dpi: int = 150):
        """
        Initialize the document converter.

        Args:
            dpi: DPI for image generation
        """
        self.dpi = dpi
        self.page_width = int(8.5 * dpi)  # 8.5 inches at specified DPI
        self.page_height = int(11 * dpi)  # 11 inches at specified DPI
        self.margin = int(0.5 * dpi)  # 0.5 inch margin

    def convert_text_to_pages(self, content: str) -> List[Tuple[bytes, str]]:
        """
        Convert plain text content to page images and text.

        Args:
            content: Plain text content

        Returns:
            List of tuples (image_bytes, page_text)
        """
        try:
            # Use a basic font
            try:
                font = ImageFont.truetype("DejaVuSansMono.ttf", 12)
            except OSError:
                font = ImageFont.load_default()

            # Calculate text area dimensions
            text_width = self.page_width - (2 * self.margin)
            text_height = self.page_height - (2 * self.margin)

            # Split content into lines and wrap long lines
            lines = []
            for line in content.split("\n"):
                if not line.strip():
                    lines.append("")
                    continue

                # Estimate characters per line based on font and width
                avg_char_width = 7  # Approximate for monospace font
                chars_per_line = text_width // avg_char_width

                if len(line) <= chars_per_line:
                    lines.append(line)
                else:
                    # Wrap long lines
                    while len(line) > chars_per_line:
                        lines.append(line[:chars_per_line])
                        line = line[chars_per_line:]
                    if line:
                        lines.append(line)

            # Calculate lines per page
            line_height = 16  # Approximate line height
            lines_per_page = text_height // line_height

            # Split into pages
            pages = []
            for i in range(0, len(lines), lines_per_page):
                page_lines = lines[i : i + lines_per_page]
                page_text = "\n".join(page_lines)

                # Create image
                img = Image.new("RGB", (self.page_width, self.page_height), "white")
                draw = ImageDraw.Draw(img)

                # Draw text
                y_pos = self.margin
                for line in page_lines:
                    draw.text((self.margin, y_pos), line, fill="black", font=font)
                    y_pos += line_height

                # Convert to bytes
                img_buffer = io.BytesIO()
                img.save(img_buffer, format="JPEG", quality=95)
                img_bytes = img_buffer.getvalue()

                pages.append((img_bytes, page_text))

            return pages if pages else [(self._create_empty_page(), "")]

        except Exception as e:
            logger.error(f"Error converting text to pages: {str(e)}")
            return [(self._create_empty_page(), content)]

    def convert_csv_to_pages(self, content: str) -> List[Tuple[bytes, str]]:
        """
        Convert CSV content to page images and text with enhanced pandas processing.

        Args:
            content: CSV content as string

        Returns:
            List of tuples (image_bytes, page_text)
        """
        try:
            import csv

            import pandas as pd

            # First try pandas for intelligent processing
            try:
                # Use pandas to read CSV with automatic type inference
                df = pd.read_csv(
                    io.StringIO(content),
                    dtype_backend="numpy_nullable",  # Better null handling
                    parse_dates=True,  # Automatic date parsing
                )

                if df.empty:
                    return [(self._create_empty_page(), "")]

                # Generate high-quality markdown using pandas
                formatted_text = self._format_csv_with_pandas(df, content)

            except Exception as pandas_error:
                logger.warning(
                    f"Pandas CSV processing failed, falling back to basic parsing: {pandas_error}"
                )
                # Fallback to basic CSV parsing
                csv_reader = csv.reader(io.StringIO(content))
                rows = list(csv_reader)

                if not rows:
                    return [(self._create_empty_page(), "")]

                # Format as table text using improved method
                formatted_text = self._format_csv_as_table(rows)

            # Convert the enhanced markdown text to clean page images
            return self._convert_markdown_to_pages(formatted_text)

        except Exception as e:
            logger.error(f"Error converting CSV to pages: {str(e)}")
            return [(self._create_empty_page(), content)]

    def convert_excel_to_pages(self, file_bytes: bytes) -> List[Tuple[bytes, str]]:
        """
        Convert Excel file to page images and text with enhanced formatting preservation.

        Args:
            file_bytes: Excel file bytes

        Returns:
            List of tuples (image_bytes, page_text)
        """
        try:
            import pandas as pd

            # Read Excel file
            with tempfile.NamedTemporaryFile(suffix=".xlsx") as tmp_file:
                tmp_file.write(file_bytes)
                tmp_file.flush()

                # Read all sheets and extract formatted data
                excel_file = pd.ExcelFile(tmp_file.name)
                formatted_elements = []

                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(tmp_file.name, sheet_name=sheet_name)

                    if df.empty:
                        continue

                    # Add sheet header element
                    formatted_elements.append(
                        {
                            "type": "sheet_header",
                            "sheet_name": sheet_name,
                            "space_before": 20,
                            "space_after": 15,
                        }
                    )

                    # Convert DataFrame to formatted table data
                    table_data = self._extract_excel_table_data(df)

                    if table_data:
                        formatted_elements.append(
                            {
                                "type": "excel_table",
                                "data": table_data,
                                "sheet_name": sheet_name,
                                "space_before": 10,
                                "space_after": 20,
                            }
                        )

                # Render formatted Excel content
                return self._render_formatted_excel_content(formatted_elements)

        except Exception as e:
            logger.error(f"Error converting Excel to pages: {str(e)}")
            return [(self._create_empty_page(), "Error reading Excel file")]

    def convert_word_to_pages(self, file_bytes: bytes) -> List[Tuple[bytes, str]]:
        """
        Convert Word document to page images and text with enhanced formatting.

        Args:
            file_bytes: Word document bytes

        Returns:
            List of tuples (image_bytes, page_text)
        """
        try:
            from docx import Document

            # Read Word document
            with tempfile.NamedTemporaryFile() as tmp_file:
                tmp_file.write(file_bytes)
                tmp_file.flush()

                doc = Document(tmp_file.name)

                # Extract formatted elements
                elements = self._extract_word_formatting(doc)

                # Render with enhanced formatting
                return self._render_formatted_word_content(elements)

        except Exception as e:
            logger.error(f"Error converting Word to pages: {str(e)}")
            return [(self._create_empty_page(), "Error reading Word document")]

    def _extract_word_formatting(self, doc) -> List[dict]:
        """Extract formatted content from Word document."""
        elements = []

        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Process paragraphs
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    elements.append({"type": "spacing", "height": 12})
                    continue

                # Determine if this is a heading
                style_name = paragraph.style.name if paragraph.style else "Normal"
                is_heading = "Heading" in style_name
                heading_level = 0

                if is_heading:
                    try:
                        heading_level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        heading_level = 1

                # Get alignment
                alignment = "left"
                if paragraph.alignment:
                    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        alignment = "center"
                    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                        alignment = "right"
                    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        alignment = "justify"

                # Extract run-level formatting
                formatted_runs = []
                for run in paragraph.runs:
                    if run.text.strip():
                        run_info = {
                            "text": run.text,
                            "bold": bool(run.bold),
                            "italic": bool(run.italic),
                            "underline": bool(run.underline),
                            "font_size": (
                                getattr(run.font.size, "pt", None)
                                if run.font.size
                                else None
                            ),
                            "font_name": run.font.name if run.font.name else None,
                        }
                        formatted_runs.append(run_info)

                if not formatted_runs:
                    formatted_runs = [
                        {
                            "text": paragraph.text,
                            "bold": False,
                            "italic": False,
                            "underline": False,
                            "font_size": None,
                            "font_name": None,
                        }
                    ]

                para_element = {
                    "type": "paragraph",
                    "text": paragraph.text,
                    "style": style_name,
                    "is_heading": is_heading,
                    "heading_level": heading_level,
                    "alignment": alignment,
                    "runs": formatted_runs,
                    "space_before": 6 if is_heading else 3,
                    "space_after": 6 if is_heading else 3,
                }
                elements.append(para_element)

            # Process tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        # Check if first row (likely header)
                        is_header = table.rows[0] == row
                        cell_info = {
                            "text": cell_text,
                            "is_header": is_header,
                            "bold": is_header,  # Headers are typically bold
                            "alignment": "center" if is_header else "left",
                        }
                        row_data.append(cell_info)
                    table_data.append(row_data)

                if table_data:
                    elements.append(
                        {
                            "type": "table",
                            "data": table_data,
                            "space_before": 12,
                            "space_after": 12,
                        }
                    )

        except Exception as e:
            logger.error(f"Error extracting Word formatting: {str(e)}")
            # Fallback to simple text extraction
            text_content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elements = [
                {
                    "type": "paragraph",
                    "text": text_content,
                    "style": "Normal",
                    "is_heading": False,
                    "heading_level": 0,
                    "alignment": "left",
                    "runs": [
                        {
                            "text": text_content,
                            "bold": False,
                            "italic": False,
                            "underline": False,
                            "font_size": None,
                            "font_name": None,
                        }
                    ],
                }
            ]

        return elements

    def _render_formatted_word_content(
        self, elements: List[dict]
    ) -> List[Tuple[bytes, str]]:
        """Render formatted Word content with enhanced typography."""
        try:
            # Load fonts
            fonts = self._load_fonts()

            # Calculate layout
            pages_content = self._calculate_word_page_layout(elements)

            # Render pages
            pages = []
            for page_elements in pages_content:
                img_bytes, text = self._render_word_page(page_elements, fonts)
                pages.append((img_bytes, text))

            return pages if pages else [(self._create_empty_page(), "")]

        except Exception as e:
            logger.error(f"Error rendering formatted Word content: {str(e)}")
            # Fallback to simple text rendering
            text_content = "\n".join(
                [elem.get("text", "") for elem in elements if elem.get("text")]
            )
            return self.convert_text_to_pages(text_content)

    def _load_fonts(self) -> dict:
        """Load available fonts with fallbacks."""
        fonts = {}

        # Font size hierarchy
        font_sizes = {
            "heading1": 24,
            "heading2": 20,
            "heading3": 18,
            "heading4": 16,
            "heading5": 14,
            "heading6": 13,
            "normal": 12,
            "small": 10,
        }

        # Try to load system fonts
        font_paths = [
            # Windows
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/calibri.ttf",
            "C:/Windows/Fonts/times.ttf",
            # macOS
            "/System/Library/Fonts/Arial.ttf",
            "/System/Library/Fonts/Times.ttc",
            "/System/Library/Fonts/Helvetica.ttc",
            # Linux
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            "/usr/share/fonts/truetype/ubuntu/Ubuntu-R.ttf",
        ]

        loaded_font = None
        for font_path in font_paths:
            try:
                if os.path.exists(font_path):
                    loaded_font = font_path
                    break
            except OSError:
                continue

        # Create font dictionary
        for name, size in font_sizes.items():
            try:
                if loaded_font:
                    fonts[name] = ImageFont.truetype(loaded_font, size)
                else:
                    fonts[name] = ImageFont.load_default()
            except (OSError, IOError):
                fonts[name] = ImageFont.load_default()

        return fonts

    def _calculate_word_page_layout(self, elements: List[dict]) -> List[List[dict]]:
        """Calculate page breaks for Word content."""
        pages = []
        current_page = []
        current_y = self.margin

        # Usable page height
        max_y = self.page_height - self.margin

        for element in elements:
            # Estimate element height
            if element["type"] == "spacing":
                element_height = element["height"]
            elif element["type"] == "paragraph":
                # Estimate based on text length and heading level
                base_height = 24 if element["is_heading"] else 16
                lines = max(1, len(element["text"]) // 80 + 1)  # Rough estimate
                element_height = (
                    base_height * lines
                    + element.get("space_before", 0)
                    + element.get("space_after", 0)
                )
            elif element["type"] == "table":
                # Estimate table height
                row_count = len(element["data"])
                element_height = (
                    row_count * 25
                    + element.get("space_before", 0)
                    + element.get("space_after", 0)
                )
            else:
                element_height = 20

            # Check if we need a new page
            if current_y + element_height > max_y and current_page:
                pages.append(current_page)
                current_page = []
                current_y = self.margin

            current_page.append(element)
            current_y += element_height

        # Add final page
        if current_page:
            pages.append(current_page)

        return pages if pages else [[]]

    def _render_word_page(self, elements: List[dict], fonts: dict) -> Tuple[bytes, str]:
        """Render a single page with enhanced formatting."""
        try:
            # Create image
            img = Image.new("RGB", (self.page_width, self.page_height), "white")
            draw = ImageDraw.Draw(img)

            # Track position and collect text
            y_pos = self.margin
            page_text = []
            text_width = self.page_width - (2 * self.margin)

            for element in elements:
                if element["type"] == "spacing":
                    y_pos += element["height"]

                elif element["type"] == "paragraph":
                    y_pos += element.get("space_before", 0)

                    # Choose font based on heading level
                    if element["is_heading"]:
                        font_key = f"heading{min(element['heading_level'], 6)}"
                    else:
                        font_key = "normal"

                    font = fonts.get(font_key, fonts["normal"])

                    # Render paragraph with formatting
                    para_height = self._render_formatted_paragraph(
                        draw, element, self.margin, y_pos, text_width, font
                    )

                    y_pos += para_height + element.get("space_after", 0)
                    page_text.append(element["text"])

                elif element["type"] == "table":
                    y_pos += element.get("space_before", 0)

                    # Render table
                    table_height = self._render_formatted_table(
                        draw, element["data"], self.margin, y_pos, text_width, fonts
                    )

                    y_pos += table_height + element.get("space_after", 0)

                    # Add table text
                    for row in element["data"]:
                        row_text = " | ".join([cell["text"] for cell in row])
                        page_text.append(row_text)

            # Convert to bytes
            img_buffer = io.BytesIO()
            img.save(img_buffer, format="JPEG", quality=95)
            img_buffer.seek(0)
            img_bytes = img_buffer.getvalue()

            return img_bytes, "\n".join(page_text)

        except Exception as e:
            logger.error(f"Error rendering Word page: {str(e)}")
            return self._create_empty_page(), "\n".join(
                [elem.get("text", "") for elem in elements]
            )

    def _render_formatted_paragraph(
        self, draw, element: dict, x: int, y: int, width: int, base_font
    ) -> int:
        """Render a paragraph with run-level formatting."""
        try:
            current_x = x
            current_y = y
            line_height = 20 if element["is_heading"] else 16

            # Handle alignment for simple case
            if element["alignment"] == "center":
                # Simple center alignment - measure total text width
                total_width = self._get_text_width(draw, element["text"], base_font)
                if total_width < width:
                    current_x = x + (width - total_width) // 2
            elif element["alignment"] == "right":
                total_width = self._get_text_width(draw, element["text"], base_font)
                if total_width < width:
                    current_x = x + width - total_width

            # Render runs with formatting
            for run in element["runs"]:
                if not run["text"]:
                    continue

                # Apply formatting effects
                if run["bold"]:
                    # Simulate bold by drawing multiple times with slight offsets
                    for offset in [(0, 0), (1, 0), (0, 1), (1, 1)]:
                        draw.text(
                            (current_x + offset[0], current_y + offset[1]),
                            run["text"],
                            fill="black",
                            font=base_font,
                        )
                else:
                    draw.text(
                        (current_x, current_y),
                        run["text"],
                        fill="black",
                        font=base_font,
                    )

                # Handle underline
                if run["underline"]:
                    text_width = self._get_text_width(draw, run["text"], base_font)
                    underline_y = current_y + line_height - 2
                    draw.line(
                        [
                            (current_x, underline_y),
                            (current_x + text_width, underline_y),
                        ],
                        fill="black",
                        width=1,
                    )

                # Move x position for next run
                run_width = self._get_text_width(draw, run["text"], base_font)
                current_x += run_width

                # Simple line wrapping
                if current_x > x + width:
                    current_x = x
                    current_y += line_height

            return max(line_height, current_y - y + line_height)

        except Exception as e:
            logger.error(f"Error rendering paragraph: {str(e)}")
            # Fallback to simple rendering
            draw.text((x, y), element["text"], fill="black", font=base_font)
            return 20

    def _render_formatted_table(
        self,
        draw,
        table_data: List[List[dict]],
        x: int,
        y: int,
        width: int,
        fonts: dict,
    ) -> int:
        """Render a table with borders and formatting."""
        try:
            if not table_data:
                return 0

            # Calculate column widths
            col_count = len(table_data[0])
            col_width = width // col_count
            row_height = 25

            current_y = y

            for row_idx, row in enumerate(table_data):
                current_x = x

                # Draw row background for headers
                if row and row[0].get("is_header", False):
                    draw.rectangle(
                        [x, current_y, x + width, current_y + row_height],
                        fill="#f0f0f0",
                        outline="#cccccc",
                    )

                for col_idx, cell in enumerate(row):
                    # Draw cell border
                    cell_rect = [
                        current_x,
                        current_y,
                        current_x + col_width,
                        current_y + row_height,
                    ]
                    draw.rectangle(cell_rect, outline="#cccccc", width=1)

                    # Choose font and formatting
                    font = fonts["normal"]
                    if cell.get("is_header", False):
                        font = fonts[
                            "normal"
                        ]  # Will be made bold by rendering multiple times

                    # Calculate text position with padding
                    text_x = current_x + 5
                    text_y = current_y + 5

                    # Handle cell alignment
                    if cell.get("alignment") == "center":
                        text_width = self._get_text_width(draw, cell["text"], font)
                        if text_width < col_width - 10:
                            text_x = current_x + (col_width - text_width) // 2
                    elif cell.get("alignment") == "right":
                        text_width = self._get_text_width(draw, cell["text"], font)
                        text_x = current_x + col_width - text_width - 5

                    # Render cell text
                    if cell.get("bold", False):
                        # Simulate bold
                        for offset in [(0, 0), (1, 0), (0, 1)]:
                            draw.text(
                                (text_x + offset[0], text_y + offset[1]),
                                cell["text"],
                                fill="black",
                                font=font,
                            )
                    else:
                        draw.text(
                            (text_x, text_y), cell["text"], fill="black", font=font
                        )

                    current_x += col_width

                current_y += row_height

            return current_y - y

        except Exception as e:
            logger.error(f"Error rendering table: {str(e)}")
            # Fallback to simple table rendering
            simple_height = len(table_data) * 20
            table_y = y
            for row in table_data:
                row_text = " | ".join([cell.get("text", "") for cell in row])
                draw.text((x, table_y), row_text, fill="black", font=fonts["normal"])
                table_y += 20
            return simple_height

    def _extract_excel_table_data(self, df) -> List[List[dict]]:
        """
        Extract Excel DataFrame into formatted table data with type-aware formatting.

        Args:
            df: pandas DataFrame

        Returns:
            List of row data with cell formatting information
        """
        try:
            import pandas as pd

            if df.empty:
                return []

            table_data = []

            # Create header row
            header_row = []
            for col_name in df.columns:
                header_row.append(
                    {
                        "text": str(col_name),
                        "is_header": True,
                        "bold": True,
                        "alignment": "center",
                        "data_type": "text",
                    }
                )
            table_data.append(header_row)

            # Process data rows
            for _, row in df.iterrows():
                data_row = []
                for col_idx, (col_name, value) in enumerate(row.items()):
                    # Handle NaN/None values
                    if pd.isna(value):
                        cell_text = ""
                        data_type = "text"
                        alignment = "left"
                    else:
                        # Determine data type and formatting
                        if pd.api.types.is_numeric_dtype(df[col_name]):
                            # Format numbers appropriately
                            if isinstance(value, float):
                                if value.is_integer():
                                    cell_text = f"{int(value):,}"
                                else:
                                    cell_text = f"{value:,.2f}"
                            else:
                                cell_text = f"{value:,}"
                            data_type = "numeric"
                            alignment = "right"
                        elif pd.api.types.is_datetime64_any_dtype(df[col_name]):
                            # Format dates
                            try:
                                cell_text = value.strftime("%Y-%m-%d")
                            except (AttributeError, ValueError, TypeError):
                                cell_text = str(value)
                            data_type = "date"
                            alignment = "center"
                        else:
                            # Text data
                            cell_text = str(value)
                            data_type = "text"
                            alignment = "left"

                            # Special handling for currency-like text
                            if (
                                cell_text.startswith("$")
                                or "€" in cell_text
                                or "£" in cell_text
                            ):
                                data_type = "currency"
                                alignment = "right"

                    data_row.append(
                        {
                            "text": cell_text,
                            "is_header": False,
                            "bold": False,
                            "alignment": alignment,
                            "data_type": data_type,
                        }
                    )

                table_data.append(data_row)

            return table_data

        except Exception as e:
            logger.error(f"Error extracting Excel table data: {str(e)}")
            # Fallback to simple conversion
            simple_data = []
            try:
                # Convert to simple string representation
                for col in df.columns:
                    simple_data.append(
                        [
                            {
                                "text": str(col),
                                "is_header": True,
                                "bold": True,
                                "alignment": "center",
                                "data_type": "text",
                            }
                        ]
                    )
                    break

                for _, row in df.iterrows():
                    row_data = []
                    for value in row:
                        row_data.append(
                            {
                                "text": str(value) if not pd.isna(value) else "",
                                "is_header": False,
                                "bold": False,
                                "alignment": "left",
                                "data_type": "text",
                            }
                        )
                    simple_data.append(row_data)
                return simple_data
            except Exception:
                return []

    def _render_formatted_excel_content(
        self, elements: List[dict]
    ) -> List[Tuple[bytes, str]]:
        """
        Render formatted Excel content as clean markdown pages.

        Args:
            elements: List of formatted Excel elements (sheet headers, tables)

        Returns:
            List of tuples (image_bytes, page_text)
        """
        try:
            # Generate enhanced markdown text for Excel content
            enhanced_text = self._generate_enhanced_excel_markdown(elements)

            # Convert the enhanced markdown text to clean page images
            return self._convert_markdown_to_pages(enhanced_text)

        except Exception as e:
            logger.error(f"Error rendering formatted Excel content: {str(e)}")
            # Fallback to simple text rendering
            text_content = []
            for elem in elements:
                if elem.get("type") == "sheet_header":
                    text_content.append(
                        f"=== Sheet: {elem.get('sheet_name', 'Unknown')} ==="
                    )
                elif elem.get("type") == "excel_table" and elem.get("data"):
                    for row in elem["data"]:
                        row_text = " | ".join([cell.get("text", "") for cell in row])
                        text_content.append(row_text)

            combined_text = "\n".join(text_content)
            return self._convert_markdown_to_pages(combined_text)

    def _get_text_width(self, draw, text: str, font) -> int:
        """Get text width using the appropriate PIL method."""
        try:
            # Try new textbbox method (PIL 8.0.0+)
            bbox = draw.textbbox((0, 0), text, font=font)
            return bbox[2] - bbox[0]
        except AttributeError:
            try:
                # Fallback to deprecated textsize method
                return draw.textsize(text, font=font)[0]
            except AttributeError:
                # Ultimate fallback - estimate based on text length
                return len(text) * 8  # Rough estimation

    def _format_csv_with_pandas(self, df, original_content: str) -> str:
        """
        Format CSV using pandas - just the clean table without metadata.

        Args:
            df: pandas DataFrame
            original_content: Original CSV content for fallback

        Returns:
            Clean markdown table formatted text
        """
        try:
            import pandas as pd

            # Format numeric columns with appropriate precision
            df_formatted = df.copy()
            for col in df_formatted.columns:
                if pd.api.types.is_numeric_dtype(df_formatted[col]):
                    if pd.api.types.is_float_dtype(df_formatted[col]):
                        # Format floats with 2 decimal places, but remove trailing zeros
                        df_formatted[col] = df_formatted[col].apply(
                            lambda x: (
                                f"{x:,.2f}".rstrip("0").rstrip(".")
                                if pd.notna(x)
                                else ""
                            )
                        )
                    else:
                        # Format integers with thousand separators
                        df_formatted[col] = df_formatted[col].apply(
                            lambda x: f"{x:,}" if pd.notna(x) else ""
                        )
                elif pd.api.types.is_datetime64_any_dtype(df_formatted[col]):
                    # Format dates consistently
                    df_formatted[col] = df_formatted[col].dt.strftime("%Y-%m-%d")

            # Generate clean markdown table only
            table_markdown = df_formatted.to_markdown(
                index=False, tablefmt="pipe", stralign="left", numalign="right"
            )

            return table_markdown

        except Exception as e:
            logger.error(f"Error in pandas CSV formatting: {str(e)}")
            # Fallback to basic CSV parsing
            import csv

            csv_reader = csv.reader(io.StringIO(original_content))
            rows = list(csv_reader)
            return self._format_csv_as_table(rows)

    def _generate_enhanced_excel_markdown(self, elements: List[dict]) -> str:
        """
        Generate clean markdown for Excel content - just the table/sheet contents without metadata.

        Args:
            elements: List of Excel elements (sheet headers, tables)

        Returns:
            Clean markdown formatted text with just table contents
        """
        try:
            import pandas as pd

            markdown_parts = []

            # Process each element
            for element in elements:
                if element.get("type") == "sheet_header":
                    sheet_name = element.get("sheet_name", "Unknown")
                    # Only add sheet name as a simple header if there are multiple sheets
                    if (
                        len(
                            [
                                elem
                                for elem in elements
                                if elem.get("type") == "sheet_header"
                            ]
                        )
                        > 1
                    ):
                        markdown_parts.append(f"## {sheet_name}")
                        markdown_parts.append("")

                elif element.get("type") == "excel_table":
                    table_data = element.get("data", [])

                    if not table_data:
                        continue

                    # Convert table data back to DataFrame for clean formatting
                    try:
                        # Extract headers and data
                        headers = [
                            cell.get("text", "")
                            for cell in table_data[0]
                            if cell.get("is_header", False)
                        ]
                        data_rows = []

                        for row in table_data[1:]:  # Skip header row
                            row_data = []
                            for cell in row:
                                cell_text = cell.get("text", "")
                                data_type = cell.get("data_type", "text")

                                # Convert back to appropriate type for formatting
                                if data_type == "numeric" and cell_text:
                                    try:
                                        # Remove commas and convert to number
                                        numeric_value = float(
                                            cell_text.replace(",", "")
                                        )
                                        if numeric_value.is_integer():
                                            row_data.append(int(numeric_value))
                                        else:
                                            row_data.append(numeric_value)
                                    except (ValueError, AttributeError):
                                        row_data.append(cell_text)
                                else:
                                    row_data.append(cell_text if cell_text else None)

                            data_rows.append(row_data)

                        if headers and data_rows:
                            # Create DataFrame for clean table formatting
                            df = pd.DataFrame(data_rows, columns=headers)

                            # Format the DataFrame for display
                            df_display = df.copy()
                            for col in df_display.columns:
                                if pd.api.types.is_numeric_dtype(df_display[col]):
                                    if pd.api.types.is_float_dtype(df_display[col]):
                                        df_display[col] = df_display[col].apply(
                                            lambda x: (
                                                f"{x:,.2f}".rstrip("0").rstrip(".")
                                                if pd.notna(x)
                                                else ""
                                            )
                                        )
                                    else:
                                        df_display[col] = df_display[col].apply(
                                            lambda x: f"{x:,}" if pd.notna(x) else ""
                                        )

                            # Generate clean markdown table only
                            table_markdown = df_display.to_markdown(
                                index=False,
                                tablefmt="pipe",
                                stralign="left",
                                numalign="right",
                            )
                            markdown_parts.append(table_markdown)
                            markdown_parts.append("")  # Add spacing between tables

                    except Exception as e:
                        logger.warning(
                            f"Error processing Excel table data for clean markdown: {str(e)}"
                        )
                        # Fallback to simple table representation
                        for row in table_data:
                            row_text = " | ".join(
                                [cell.get("text", "") for cell in row]
                            )
                            markdown_parts.append(f"| {row_text} |")
                        markdown_parts.append("")

            return "\n".join(markdown_parts)

        except Exception as e:
            logger.error(f"Error generating clean Excel markdown: {str(e)}")
            # Fallback to simple representation
            fallback_parts = []
            for elem in elements:
                if elem.get("type") == "sheet_header":
                    sheet_name = elem.get("sheet_name", "Unknown")
                    # Only add sheet name if there are multiple sheets
                    if (
                        len([e for e in elements if e.get("type") == "sheet_header"])
                        > 1
                    ):
                        fallback_parts.append(f"## {sheet_name}")
                        fallback_parts.append("")
                elif elem.get("type") == "excel_table" and elem.get("data"):
                    for row in elem["data"]:
                        row_text = " | ".join([cell.get("text", "") for cell in row])
                        fallback_parts.append(f"| {row_text} |")
                    fallback_parts.append("")
            return "\n".join(fallback_parts)

    def _format_csv_as_table(self, rows: List[List[str]]) -> str:
        """Format CSV rows as a readable table in proper markdown format."""
        if not rows:
            return ""

        # Calculate column widths without truncation
        col_widths = []
        for col_idx in range(len(rows[0])):
            max_width = 0
            for row in rows:
                if col_idx < len(row):
                    max_width = max(max_width, len(str(row[col_idx])))
            col_widths.append(max_width)  # No cap on characters

        # Format rows as markdown table
        formatted_rows = []
        for row_idx, row in enumerate(rows):
            formatted_cells = []
            for col_idx, cell in enumerate(row):
                if col_idx < len(col_widths):
                    cell_str = str(cell)  # Include all text without truncation
                    formatted_cells.append(cell_str)

            # Proper markdown table format with leading and trailing pipes
            formatted_row = "| " + " | ".join(formatted_cells) + " |"
            formatted_rows.append(formatted_row)

            # Add markdown separator after header
            if row_idx == 0 and len(rows) > 1:
                # Create separator line with proper markdown syntax
                separators = []
                for width in col_widths:
                    # Use at least 3 dashes for markdown compliance
                    separator = "-" * max(3, width)
                    separators.append(separator)
                separator_row = "| " + " | ".join(separators) + " |"
                formatted_rows.append(separator_row)

        return "\n".join(formatted_rows)

    def _convert_markdown_to_pages(
        self, markdown_content: str
    ) -> List[Tuple[bytes, str]]:
        """
        Convert markdown content to clean page images with proper formatting.
        Returns original markdown as page_text to preserve proper markdown syntax.

        Args:
            markdown_content: Markdown formatted text

        Returns:
            List of tuples (image_bytes, page_text)
        """
        try:
            # Use a monospace font for better markdown rendering
            try:
                font_normal = ImageFont.truetype("DejaVuSansMono.ttf", 12)
                font_bold = ImageFont.truetype("DejaVuSansMono-Bold.ttf", 12)
                font_heading = ImageFont.truetype("DejaVuSansMono-Bold.ttf", 16)
            except OSError:
                font_normal = ImageFont.load_default()
                font_bold = ImageFont.load_default()
                font_heading = ImageFont.load_default()

            # Calculate text area dimensions
            text_width = self.page_width - (2 * self.margin)
            text_height = self.page_height - (2 * self.margin)

            # Calculate lines per page with better spacing
            line_height = 18  # Slightly more space for better readability
            lines_per_page = text_height // line_height

            # Split the original markdown into pages while preserving table structure
            original_lines = markdown_content.split("\n")

            # Find table headers and separators in the original markdown
            table_info = self._analyze_table_structure(original_lines)

            pages = []
            original_line_idx = 0

            while original_line_idx < len(original_lines):
                # Get a chunk of original lines for this page
                page_original_lines = original_lines[
                    original_line_idx : original_line_idx + lines_per_page
                ]

                # Check if this page starts in the middle of a table
                page_text_lines = self._ensure_table_headers(
                    page_original_lines, table_info, original_line_idx
                )

                # Create the page text from processed markdown
                page_text = "\n".join(page_text_lines)

                # Create image with simple but clean formatting
                img = Image.new("RGB", (self.page_width, self.page_height), "white")
                draw = ImageDraw.Draw(img)

                # Render with simple text formatting (fast and preserves all content)
                y_pos = self.margin

                for line in page_text_lines:
                    if y_pos + line_height > self.page_height - self.margin:
                        break  # Page is full

                    # Simple formatting based on content
                    if line.startswith("#"):
                        # Heading - use bold font and remove markdown syntax
                        text = line.lstrip("#").strip()
                        font = font_heading
                        color = "#2c3e50"
                    elif line.startswith("- ") or line.startswith("* "):
                        # List item - add bullet and indent
                        text = "• " + line[2:].strip()
                        font = font_normal
                        color = "black"
                        x_pos = self.margin + 20
                    elif "**" in line:
                        # Bold text - remove markdown and use bold font
                        text = line.replace("**", "")
                        font = font_bold
                        color = "black"
                    else:
                        # Regular text
                        text = line
                        font = font_normal
                        color = "black"

                    # Default x position
                    if not line.startswith("- ") and not line.startswith("* "):
                        x_pos = self.margin

                    # Handle long lines by wrapping
                    wrapped_lines = self._wrap_text_to_width(
                        text, font, text_width - (x_pos - self.margin), draw
                    )

                    for wrapped_line in wrapped_lines:
                        if y_pos + line_height > self.page_height - self.margin:
                            break  # Page is full

                        # Draw the text
                        draw.text((x_pos, y_pos), wrapped_line, fill=color, font=font)
                        y_pos += line_height

                    # Add small spacing after headings
                    if line.startswith("#"):
                        y_pos += 6

                # Convert to bytes
                img_buffer = io.BytesIO()
                img.save(img_buffer, format="JPEG", quality=95)
                img_bytes = img_buffer.getvalue()

                pages.append((img_bytes, page_text))
                original_line_idx += len(page_original_lines)

            return pages if pages else [(self._create_empty_page(), markdown_content)]

        except Exception as e:
            logger.error(f"Error converting markdown to pages: {str(e)}")
            # Fallback to basic text conversion
            return self.convert_text_to_pages(markdown_content)

    def _wrap_text_to_width(self, text: str, font, max_width: int, draw) -> List[str]:
        """
        Wrap text to fit within specified width.

        Args:
            text: Text to wrap
            font: Font to use for measurement
            max_width: Maximum width in pixels
            draw: PIL ImageDraw object for text measurement

        Returns:
            List of wrapped text lines
        """
        if not text.strip():
            return [text]

        words = text.split()
        if not words:
            return [text]

        lines = []
        current_line = []

        for word in words:
            # Test if adding this word would exceed width
            test_line = " ".join(current_line + [word])
            text_width = self._get_text_width(draw, test_line, font)

            if text_width <= max_width or not current_line:
                current_line.append(word)
            else:
                # Start new line
                if current_line:
                    lines.append(" ".join(current_line))
                current_line = [word]

        # Add remaining words
        if current_line:
            lines.append(" ".join(current_line))

        return lines if lines else [text]

    def _analyze_table_structure(self, lines: List[str]) -> dict:
        """
        Analyze the markdown content to find table headers and their positions.

        Args:
            lines: List of markdown lines

        Returns:
            Dictionary with table structure information
        """
        table_info = {
            "headers": [],  # List of (line_idx, header_line, separator_line)
            "table_ranges": [],  # List of (start_idx, end_idx) for each table
        }

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Check if this is a table header (line with | followed by separator)
            if (
                line.startswith("|")
                and "|" in line[1:]
                and i + 1 < len(lines)
                and "---" in lines[i + 1]
                and "|" in lines[i + 1]
            ):
                header_line = lines[i]
                separator_line = lines[i + 1]

                # Find the end of this table
                table_start = i
                table_end = i + 1

                # Look for more table rows
                j = i + 2
                while j < len(lines):
                    next_line = lines[j].strip()
                    if next_line.startswith("|") and "|" in next_line[1:]:
                        table_end = j
                        j += 1
                    else:
                        break

                table_info["headers"].append((i, header_line, separator_line))
                table_info["table_ranges"].append((table_start, table_end))

                i = table_end + 1
            else:
                i += 1

        return table_info

    def _ensure_table_headers(
        self, page_lines: List[str], table_info: dict, start_line_idx: int
    ) -> List[str]:
        """
        Ensure that if a page starts in the middle of a table, it includes the table header.

        Args:
            page_lines: Lines for this page
            table_info: Table structure information
            start_line_idx: Starting line index in the original document

        Returns:
            Modified page lines with table headers if needed
        """
        if not page_lines or not table_info["table_ranges"]:
            return page_lines

        # Check if this page starts in the middle of a table
        for table_start, table_end in table_info["table_ranges"]:
            if table_start < start_line_idx <= table_end:
                # This page starts in the middle of a table
                # Find the corresponding header
                for header_idx, header_line, separator_line in table_info["headers"]:
                    if table_start <= header_idx <= table_end:
                        # Add the header and separator to the beginning of the page
                        result_lines = [header_line, separator_line] + page_lines
                        return result_lines

        return page_lines

    def _create_empty_page(self) -> bytes:
        """Create an empty white page image."""
        try:
            img = Image.new("RGB", (self.page_width, self.page_height), "white")
            img_buffer = io.BytesIO()
            img.save(img_buffer, format="JPEG", quality=95)
            img_buffer.seek(0)  # Reset buffer position to beginning
            img_bytes = img_buffer.getvalue()

            # Ensure we actually got bytes
            if len(img_bytes) > 0:
                return img_bytes
            else:
                logger.warning("Image save produced 0 bytes, creating minimal JPEG")
                # Create a minimal 1x1 white JPEG as fallback
                minimal_img = Image.new("RGB", (1, 1), "white")
                minimal_buffer = io.BytesIO()
                minimal_img.save(minimal_buffer, format="JPEG", quality=95)
                minimal_buffer.seek(0)
                minimal_bytes = minimal_buffer.getvalue()
                logger.warning(f"Minimal JPEG created with {len(minimal_bytes)} bytes")
                if len(minimal_bytes) > 0:
                    return minimal_bytes
                else:
                    # Fall through to hardcoded JPEG
                    logger.error(
                        "Even minimal JPEG creation failed, using hardcoded bytes"
                    )
                    # Continue to hardcoded fallback below

        except Exception as e:
            logger.error(f"Error creating empty page: {str(e)}")
            # Fall through to hardcoded minimal JPEG

        # Return a hardcoded minimal valid 1x1 white JPEG
        logger.warning("Using hardcoded minimal JPEG")
        return b"\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01\x01\x01\x00H\x00H\x00\x00\xff\xdb\x00C\x00\x08\x06\x06\x07\x06\x05\x08\x07\x07\x07\t\t\x08\n\x0c\x14\r\x0c\x0b\x0b\x0c\x19\x12\x13\x0f\x14\x1d\x1a\x1f\x1e\x1d\x1a\x1c\x1c $.' \",#\x1c\x1c(7),01444\x1f'9=82<.342\xff\xc0\x00\x11\x08\x00\x01\x00\x01\x01\x01\x11\x00\x02\x11\x01\x03\x11\x01\xff\xc4\x00\x14\x00\x01\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x08\xff\xc4\x00\x14\x10\x01\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\xff\xda\x00\x0c\x03\x01\x00\x02\x11\x03\x11\x00\x3f\x00\x80\xff\xd9"
